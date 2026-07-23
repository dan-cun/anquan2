from __future__ import annotations

import argparse
import json
import math
import tempfile
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


PAGE_WIDTH_DXA = 11906
CONTENT_WIDTH_DXA = 9866
FONT_NAME = "Microsoft YaHei"
NAVY = "13263D"
TEAL = "16697A"
INK = "1F2937"
MUTED = "5B677A"
PALE = "EAF3F5"
PALE_BLUE = "E8EEF5"
PALE_RED = "FDECEC"
WHITE = "FFFFFF"
GRID = "CBD5E1"

COMPONENT_MAX = {
    "goal_achievement": 60,
    "evidence_accuracy": 15,
    "reproducibility": 10,
    "decision_log": 10,
    "safety_and_cleanup": 5,
}
COMPONENT_LABELS = {
    "goal_achievement": "目标达成",
    "evidence_accuracy": "证据准确",
    "reproducibility": "可复现性",
    "decision_log": "决策日志",
    "safety_and_cleanup": "安全与清理",
}


class ReportError(RuntimeError):
    pass


@dataclass(frozen=True)
class ReportData:
    evaluation_path: Path
    batch_root: Path
    evaluation: dict[str, Any]
    cases: list[dict[str, Any]]
    observed_targets: frozenset[str]
    preview: bool
    target_name: str
    provider: str
    model: str
    total_tokens: int
    request_count: int
    ledger_valid_count: int
    cleanup_verified_count: int
    final_answer_count: int
    evidence_count: int
    finding_count: int
    status_counts: Counter[str]


def _json(path: Path) -> dict[str, Any]:
    if not path.is_file():
        raise ReportError(f"Required JSON file is missing: {path}")
    payload = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ReportError(f"Expected a JSON object: {path}")
    return payload


def _as_list(value: Any) -> list[Any]:
    return value if isinstance(value, list) else []


def _safe_int(value: Any) -> int:
    try:
        return int(value or 0)
    except (TypeError, ValueError):
        return 0


def _safe_float(value: Any) -> float:
    try:
        result = float(value or 0)
    except (TypeError, ValueError):
        return 0.0
    return result if math.isfinite(result) else 0.0


def _find_round_case(batch_root: Path, case_id: str) -> tuple[dict[str, Any], dict[str, Any]]:
    case_root = batch_root / "round-1" / case_id
    result_path = case_root / "result.json"
    environment_path = case_root / "environment.json"
    result = _json(result_path) if result_path.is_file() else {}
    environment = _json(environment_path) if environment_path.is_file() else {}
    return result, environment


def load_report_data(
    evaluation_path: Path,
    *,
    target_name: str = "PentAGI",
    preview: bool = False,
) -> ReportData:
    evaluation_path = evaluation_path.resolve()
    evaluation = _json(evaluation_path)
    scores = _as_list(evaluation.get("scores"))
    if not scores:
        raise ReportError("evaluation.json contains no score records")
    batch_root = evaluation_path.parent
    cases: list[dict[str, Any]] = []
    targets: set[str] = set()
    providers: set[str] = set()
    models: set[str] = set()
    status_counts: Counter[str] = Counter()
    ledger_valid_count = 0
    cleanup_verified_count = 0
    final_answer_count = 0
    evidence_count = 0
    finding_count = 0
    request_count = 0
    total_tokens = 0

    for score in scores:
        if not isinstance(score, dict) or not str(score.get("case_id") or ""):
            raise ReportError("Every score record must contain case_id")
        case_id = str(score["case_id"])
        result, environment = _find_round_case(batch_root, case_id)
        summary = result.get("summary") if isinstance(result.get("summary"), dict) else {}
        report = result.get("report") if isinstance(result.get("report"), dict) else {}
        usage = result.get("usage") if isinstance(result.get("usage"), dict) else {}
        target = str(environment.get("target") or summary.get("target") or "").strip().lower()
        if target:
            targets.add(target)
        provider = str(environment.get("provider") or "").strip()
        model = str(environment.get("model") or "").strip()
        if provider:
            providers.add(provider)
        if model:
            models.add(model)
        runtime_status = str(score.get("runtime_status") or result.get("status") or "unknown")
        status_counts[runtime_status] += 1
        ledger_valid_count += int(bool(result.get("ledger_chain_valid")))
        cleanup_verified_count += int((batch_root / "round-1" / case_id / "cleanup-receipt.json").is_file())
        final_answer_count += int(bool(str(report.get("final_answer") or "").strip()))
        evidence_count += len(_as_list(report.get("evidence")))
        finding_count += len(_as_list(report.get("findings")))
        request_count += _safe_int(usage.get("request_count"))
        total_tokens += _safe_int(usage.get("total_tokens"))
        cases.append(
            {
                **score,
                "runtime_status": runtime_status,
                "result": result,
                "environment": environment,
            }
        )

    batch_summary_path = batch_root / "batch-summary.json"
    if batch_summary_path.is_file():
        batch_summary = _json(batch_summary_path)
        total_usage = batch_summary.get("total_usage")
        if isinstance(total_usage, dict):
            request_count = _safe_int(total_usage.get("request_count")) or request_count
            total_tokens = _safe_int(total_usage.get("total_tokens")) or total_tokens

    required_target = target_name.strip().lower()
    is_target = bool(targets) and targets == {required_target}
    if not is_target and not preview:
        observed = ", ".join(sorted(targets)) or "unlabeled"
        raise ReportError(
            f"Refusing to label this evaluation as {target_name}: observed target is {observed}. "
            "Use --preview only for structural testing."
        )

    return ReportData(
        evaluation_path=evaluation_path,
        batch_root=batch_root,
        evaluation=evaluation,
        cases=cases,
        observed_targets=frozenset(targets),
        preview=not is_target,
        target_name=target_name,
        provider=", ".join(sorted(providers)) or "unavailable",
        model=", ".join(sorted(models)) or "unavailable",
        total_tokens=total_tokens,
        request_count=request_count,
        ledger_valid_count=ledger_valid_count,
        cleanup_verified_count=cleanup_verified_count,
        final_answer_count=final_answer_count,
        evidence_count=evidence_count,
        finding_count=finding_count,
        status_counts=status_counts,
    )


def _font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
    ]
    for candidate in candidates:
        if candidate.is_file():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def _draw_bar_chart(
    rows: Iterable[tuple[str, float]],
    *,
    title: str,
    output: Path,
    value_format: str = "percent",
) -> None:
    materialized = list(rows)
    width = 1500
    row_height = 82
    height = 150 + max(1, len(materialized)) * row_height
    image = Image.new("RGB", (width, height), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    draw.text((56, 34), title, fill="#13263D", font=_font(36, bold=True))
    label_font = _font(25)
    value_font = _font(24, bold=True)
    left = 365
    right = width - 120
    track_width = right - left
    for index, (label, raw_value) in enumerate(materialized):
        y = 115 + index * row_height
        value = max(0.0, min(1.0, _safe_float(raw_value)))
        draw.text((56, y + 10), label, fill="#1F2937", font=label_font)
        draw.rounded_rectangle((left, y + 12, right, y + 43), radius=15, fill="#E8EEF5")
        fill_right = left + int(track_width * value)
        if fill_right > left:
            draw.rounded_rectangle((left, y + 12, fill_right, y + 43), radius=15, fill="#16697A")
        text = f"{value * 100:.1f}%" if value_format == "percent" else f"{raw_value:.1f}"
        draw.text((right + 18, y + 7), text, fill="#13263D", font=value_font)
    image.save(output)


def _set_cell_fill(cell: Any, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def _set_cell_margins(cell: Any, *, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Any, widths: list[int], *, indent: int = 120) -> None:
    if len(widths) != len(table.columns) or sum(widths) != CONTENT_WIDTH_DXA:
        raise ReportError("Table widths must match the column count and total content width")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _repeat_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def _set_run_font(run: Any, *, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def _style_paragraph(paragraph: Any, *, size: float = 9.0, bold: bool = False, color: str = INK) -> None:
    for run in paragraph.runs:
        _set_run_font(run, size=size, bold=bold, color=color)


def _clear_body(document: DocumentType) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _clear_container(container: Any) -> None:
    element = container._element
    for child in list(element):
        element.remove(child)
    element.append(OxmlElement("w:p"))


def _add_page_field(paragraph: Any) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, text, end):
        run._r.append(node)
    _set_run_font(run, size=8, color=MUTED)


def _configure_header_footer(document: DocumentType, preview: bool) -> None:
    for section in document.sections:
        _clear_container(section.header)
        table = section.header.add_table(rows=1, cols=2, width=Inches(6.85))
        _set_table_geometry(table, [4933, 4933], indent=120)
        labels = ("PENTAGI / BENCHMARK", "STRUCTURAL PREVIEW" if preview else "DETERMINISTIC SCORE REPORT")
        for index, label in enumerate(labels):
            cell = table.cell(0, index)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run(label)
            _set_run_font(run, size=7.5, bold=True, color=TEAL)
        _clear_container(section.footer)
        paragraph = section.footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run("PentAGI benchmark  ·  ")
        _set_run_font(run, size=8, color=MUTED)
        _add_page_field(paragraph)


def _add_title_block(document: DocumentType, data: ReportData) -> None:
    table = document.add_table(rows=1, cols=1)
    _set_table_geometry(table, [CONTENT_WIDTH_DXA], indent=120)
    cell = table.cell(0, 0)
    _set_cell_fill(cell, NAVY)
    _set_cell_margins(cell, top=360, start=360, bottom=360, end=360)
    eyebrow = cell.paragraphs[0]
    run = eyebrow.add_run("PENTAGI / FUSED-12 BENCHMARK")
    _set_run_font(run, size=9, bold=True, color="8FD3DC")
    title = cell.add_paragraph()
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("PentAGI 跑分结果与执行闭环诊断报告")
    _set_run_font(run, size=25, bold=True, color=WHITE)
    subtitle = cell.add_paragraph()
    run = subtitle.add_run("确定性评分 · 八板块能力画像 · 逐题证据与改进门禁")
    _set_run_font(run, size=12, bold=True, color="D8EEF1")
    sample = cell.add_paragraph()
    run = sample.add_run(
        f"{data.evaluation.get('scored_case_count', len(data.cases))} 题 · "
        f"{len(data.evaluation.get('category_score_rates') or {})} 板块 · "
        f"实验 {data.evaluation.get('experiment_id', 'unknown')}"
    )
    _set_run_font(run, size=9, color="C6D4E1")
    if data.preview:
        warning = cell.add_paragraph()
        run = warning.add_run("结构验证预览：当前数据并非 PentAGI，不得作为 PentAGI 成绩发布")
        _set_run_font(run, size=9.5, bold=True, color="FFD6D6")
    document.add_paragraph()


def _add_callout(document: DocumentType, label: str, text: str, *, fill: str = PALE) -> None:
    table = document.add_table(rows=1, cols=1)
    _set_table_geometry(table, [CONTENT_WIDTH_DXA], indent=120)
    cell = table.cell(0, 0)
    _set_cell_fill(cell, fill)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(f"{label}  ")
    _set_run_font(run, size=9.5, bold=True, color=TEAL)
    run = paragraph.add_run(text)
    _set_run_font(run, size=9.5, color=INK)
    document.add_paragraph()


def _metric_text(data: ReportData) -> list[tuple[str, str, str]]:
    evaluation = data.evaluation
    rate = _safe_float(evaluation.get("equal_weight_category_score"))
    return [
        (f"{rate * 100:.1f}%", "八板块等权综合分", "确定性 evaluator"),
        (f"{evaluation.get('scored_case_count', len(data.cases))}/{evaluation.get('expected_case_count', len(data.cases))}", "题目覆盖", str(evaluation.get("report_status") or "UNKNOWN")),
        (str(evaluation.get("task_goal_success_count", 0)), "目标达成题数", f"最终答案 {data.final_answer_count}"),
        (str(data.evidence_count), "Evidence 总数", f"Finding {data.finding_count}"),
        (f"{data.total_tokens:,}", "总 Token", f"Provider 请求 {data.request_count}"),
        (f"{data.ledger_valid_count}/{len(data.cases)}", "可验证决策链", f"清理凭据 {data.cleanup_verified_count}"),
    ]


def _add_kpis(document: DocumentType, data: ReportData) -> None:
    table = document.add_table(rows=2, cols=3)
    _set_table_geometry(table, [3288, 3289, 3289], indent=120)
    for index, (value, label, note) in enumerate(_metric_text(data)):
        cell = table.cell(index // 3, index % 3)
        _set_cell_fill(cell, "F6F9FB" if index % 2 == 0 else PALE_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        _set_run_font(run, size=17, bold=True, color=NAVY)
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label)
        _set_run_font(run, size=8.5, bold=True, color=TEAL)
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(note)
        _set_run_font(run, size=7.2, color=MUTED)
    document.add_paragraph()


def _add_metadata(document: DocumentType, data: ReportData) -> None:
    evaluated_at = str(data.evaluation.get("evaluated_at") or "unavailable")
    rows = [
        ("报告性质", "PentAGI 正式确定性跑分报告" if not data.preview else "非 PentAGI 数据的结构验证预览"),
        ("证据范围", "evaluation.json、逐题 result/environment、batch-summary 与哈希清单"),
        ("模型配置", f"{data.provider} / {data.model}"),
        ("生成时间", evaluated_at),
        ("保密边界", "不导出私有答案、API Token、模型密钥或 evaluator 内部数据"),
    ]
    table = document.add_table(rows=len(rows), cols=2)
    _set_table_geometry(table, [1900, 7966], indent=120)
    for row, (label, value) in zip(table.rows, rows):
        _set_cell_fill(row.cells[0], PALE_BLUE)
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _style_paragraph(row.cells[0].paragraphs[0], size=8.5, bold=True, color=TEAL)
        _style_paragraph(row.cells[1].paragraphs[0], size=8.5, color=INK)


def _add_heading(document: DocumentType, text: str, level: int = 1) -> Any:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def _add_bullets(document: DocumentType, items: Iterable[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)
        _style_paragraph(paragraph, size=9.2, color=INK)
        paragraph.paragraph_format.space_after = Pt(2.5)


def _add_table(
    document: DocumentType,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
    *,
    small: bool = False,
) -> Any:
    table = document.add_table(rows=1, cols=len(headers))
    _set_table_geometry(table, widths, indent=120)
    header = table.rows[0]
    _repeat_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        _set_cell_fill(cell, NAVY)
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style_paragraph(paragraph, size=7.5 if small else 8.2, bold=True, color=WHITE)
    for row_values in rows:
        row = table.add_row()
        for index, text in enumerate(row_values):
            cell = row.cells[index]
            cell.text = text
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if len(table.rows) % 2 == 1:
                _set_cell_fill(cell, "F7FAFC")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index != 1 else WD_ALIGN_PARAGRAPH.LEFT
            _style_paragraph(paragraph, size=7.3 if small else 8.0, color=INK)
    document.add_paragraph()
    return table


def _diagnostic_findings(data: ReportData) -> list[tuple[str, str, str]]:
    evaluation = data.evaluation
    findings: list[tuple[str, str, str]] = []
    if _safe_int(evaluation.get("task_goal_success_count")) == 0:
        findings.append(("P0", "任务目标未达成", "没有题目通过确定性目标判定；先检查最终答案合同和执行闭环。"))
    if data.final_answer_count == 0:
        findings.append(("P0", "结构化最终答案缺失", "逐题报告没有可评分 final_answer，精确匹配题无法获得目标分。"))
    if data.evidence_count == 0:
        findings.append(("P0", "证据链为空", "报告未导出 Evidence，证据准确与可复现分无法成立。"))
    incomplete = sum(count for status, count in data.status_counts.items() if status not in {"completed", "finished"})
    if incomplete:
        findings.append(("P0", "运行终态质量不足", f"{incomplete}/{len(data.cases)} 题未以 completed/finished 形成可评分结果。"))
    if _as_list(evaluation.get("manual_review_cases")):
        findings.append(("P1", "专用 evaluator 未闭合", "BOTS、BountyBench 或 AgentDojo 中仍有人工复核项，报告只能作为暂定成绩。"))
    if data.cleanup_verified_count < len(data.cases):
        findings.append(("P1", "清理证明不完整", f"仅 {data.cleanup_verified_count}/{len(data.cases)} 题存在清理凭据，安全与副作用分受限。"))
    if data.ledger_valid_count < len(data.cases):
        findings.append(("P1", "决策链不可完全验证", f"仅 {data.ledger_valid_count}/{len(data.cases)} 题具有可验证账本链。"))
    if not findings:
        findings.append(("P2", "未发现结构性阻断", "结果合同、证据和评分覆盖均完整，下一步应比较质量、成本和稳定性。"))
    return findings


def _conclusion(data: ReportData) -> str:
    evaluation = data.evaluation
    score = _safe_float(evaluation.get("equal_weight_category_score")) * 100
    goals = _safe_int(evaluation.get("task_goal_success_count"))
    status = str(evaluation.get("report_status") or "UNKNOWN")
    label = "结构验证数据" if data.preview else "PentAGI"
    return (
        f"{label} 在固定 {len(data.cases)} 题上的八板块等权得分为 {score:.2f}%，"
        f"确定性目标达成 {goals} 题，报告状态为 {status}。"
        "分数只来自 evaluator；诊断文字不参与计分。"
    )


def _comparison_summary(path: Path | None) -> tuple[str, float] | None:
    if path is None:
        return None
    payload = _json(path.resolve())
    label = str(payload.get("experiment_id") or path.stem)
    value = _safe_float(payload.get("equal_weight_category_score"))
    return label, value


def _build_document(
    data: ReportData,
    *,
    template_docx: Path,
    output_path: Path,
    comparison: tuple[str, float] | None,
) -> None:
    document = Document(template_docx.resolve())
    _clear_body(document)
    _configure_header_footer(document, data.preview)
    styles = document.styles
    for style_name in ("Normal", "Heading 1", "Heading 2", "Small Note"):
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_NAME)

    _add_title_block(document, data)
    _add_callout(document, "一句话结论", _conclusion(data), fill=PALE)
    _add_kpis(document, data)
    _add_metadata(document, data)
    document.add_page_break()

    _add_heading(document, "执行摘要")
    score_rate = _safe_float(data.evaluation.get("equal_weight_category_score"))
    bullets = [
        f"综合分：{score_rate * 100:.2f}%，原始总分 {data.evaluation.get('total_score', 0)}/{data.evaluation.get('total_max_score', 0)}。",
        f"覆盖：{data.evaluation.get('scored_case_count', len(data.cases))}/{data.evaluation.get('expected_case_count', len(data.cases))}；状态 {data.evaluation.get('report_status', 'UNKNOWN')}。",
        f"任务产物：final_answer {data.final_answer_count}，Evidence {data.evidence_count}，Finding {data.finding_count}。",
        f"执行状态：{', '.join(f'{name}={count}' for name, count in sorted(data.status_counts.items()))}。",
        f"资源：Provider 请求 {data.request_count}，总 Token {data.total_tokens:,}。",
    ]
    _add_bullets(document, bullets)
    verdict = "当前结果尚未形成可发布的 PentAGI 正式成绩。" if data.preview or not data.evaluation.get("complete") else "固定题集已完整评分；仍需按报告状态处理人工复核项。"
    _add_callout(document, "最终裁决", verdict, fill=PALE_RED if data.preview else PALE_BLUE)

    _add_heading(document, "1. 证据范围与评分方法")
    _add_heading(document, "1.1 数据边界", level=2)
    document.add_paragraph(
        "报告读取评分器导出的 evaluation.json、逐题 result.json/environment.json、批次摘要和哈希清单。"
        "私有答案仅在确定性 evaluator 内部用于匹配，不写入报告；模型不参与分数计算。"
    )
    _add_heading(document, "1.2 评分口径", level=2)
    _add_bullets(
        document,
        [
            "单题满分 100：目标达成 60、证据准确 15、步骤可复现 10、决策日志 10、安全与清理 5。",
            "失败、超时和空答案题保留在固定 12 题分母中。",
            "总体能力分采用八板块得分率等权平均，避免题量差异改变权重。",
            "缺少专用 oracle 的题标记 MANUAL_REVIEW_REQUIRED，成绩为暂定值。",
        ],
    )

    with tempfile.TemporaryDirectory(prefix="pentagi-report-") as temporary:
        assets = Path(temporary)
        categories = data.evaluation.get("category_score_rates") or {}
        category_chart = assets / "category-scores.png"
        _draw_bar_chart(sorted((str(k), _safe_float(v)) for k, v in categories.items()), title="八板块得分率", output=category_chart)
        case_chart = assets / "case-scores.png"
        _draw_bar_chart(
            [(str(item.get("case_id")), _safe_float(item.get("score")) / max(1, _safe_float(item.get("max_score")) or 100)) for item in data.cases],
            title="逐题得分率",
            output=case_chart,
        )

        _add_heading(document, "2. 跑分结果")
        _add_heading(document, "2.1 总体结果", level=2)
        summary_rows = [
            ["八板块等权综合分", f"{score_rate * 100:.2f}%", "主报告指标"],
            ["原始总分率", f"{_safe_float(data.evaluation.get('raw_score_rate')) * 100:.2f}%", f"{data.evaluation.get('total_score', 0)}/{data.evaluation.get('total_max_score', 0)}"],
            ["任务目标达成", str(data.evaluation.get("task_goal_success_count", 0)), f"共 {len(data.cases)} 题"],
            ["运行完成", str(data.evaluation.get("runtime_completed_count", 0)), "系统终态指标，不等同于答对"],
            ["人工复核", str(len(_as_list(data.evaluation.get("manual_review_cases")))), ", ".join(_as_list(data.evaluation.get("manual_review_cases"))) or "无"],
        ]
        _add_table(document, ["指标", "结果", "解释"], summary_rows, [3100, 1800, 4966])
        if comparison:
            label, value = comparison
            delta = score_rate - value
            _add_callout(document, "对照", f"相对 {label}（{value * 100:.2f}%），本轮变化 {delta * 100:+.2f} 个百分点。")

        _add_heading(document, "2.2 八板块结果", level=2)
        document.add_picture(str(category_chart), width=Inches(6.6))
        caption = document.add_paragraph(style="Small Note")
        caption.add_run("图 1｜各板块得分率；总体能力分为八板块等权平均。")
        category_rows = [[str(name), f"{_safe_float(value) * 100:.2f}%"] for name, value in sorted(categories.items())]
        _add_table(document, ["板块", "得分率"], category_rows, [6900, 2966])

        _add_heading(document, "2.3 逐题结果", level=2)
        document.add_picture(str(case_chart), width=Inches(6.6))
        caption = document.add_paragraph(style="Small Note")
        caption.add_run("图 2｜固定 12 题逐题得分率；失败题仍保留在分母。")
        case_rows: list[list[str]] = []
        for item in data.cases:
            case_rows.append(
                [
                    str(item.get("case_id")),
                    str(item.get("category")),
                    str(item.get("runtime_status")),
                    f"{_safe_float(item.get('score')):.0f}/{_safe_float(item.get('max_score')) or 100:.0f}",
                    "是" if item.get("goal_met") else "否",
                    str(item.get("evidence_count", 0)),
                    str(item.get("score_status")),
                ]
            )
        _add_table(
            document,
            ["题目", "板块", "状态", "得分", "目标", "证据", "判定"],
            case_rows,
            [1250, 1850, 1100, 900, 700, 700, 3366],
            small=True,
        )

    _add_heading(document, "3. 执行质量与成本")
    execution_rows = [
        ["Provider 请求", f"{data.request_count:,}", "模型调用数量"],
        ["总 Token", f"{data.total_tokens:,}", "仅在导出可用时统计"],
        ["结构化最终答案", str(data.final_answer_count), f"共 {len(data.cases)} 题"],
        ["Evidence / Finding", f"{data.evidence_count} / {data.finding_count}", "任务级可验证产物"],
        ["账本链有效", f"{data.ledger_valid_count}/{len(data.cases)}", "PentAGI 适配器当前可能无法证明链式账本"],
        ["清理凭据", f"{data.cleanup_verified_count}/{len(data.cases)}", "资源与 Flow 清理证明"],
    ]
    _add_table(document, ["指标", "结果", "含义"], execution_rows, [3100, 1900, 4866])

    _add_heading(document, "4. 确定性诊断")
    findings = _diagnostic_findings(data)
    finding_rows = [[priority, title, evidence] for priority, title, evidence in findings]
    _add_table(document, ["优先级", "诊断", "证据与影响"], finding_rows, [1200, 2600, 6066])
    document.add_paragraph(
        "这些诊断由评分产物的结构化字段触发，用于解释分数和安排复测；它们不修改 evaluator 产生的任何分值。"
    )

    _add_heading(document, "5. 修复路线与复测门禁")
    roadmap = [
        ["P0", "结果合同", "每个终态任务必须输出独立 final_answer、evidence、findings、reproduction_steps。", "静态单题目标分可由 evaluator 识别"],
        ["P0", "执行闭环", "修复任务完成判定；无答案或无证据只能 partial/unsupported。", "空证据不得 completed"],
        ["P1", "能力匹配", "按 Web/Pwn/Reverse/Crypto/DFIR/Repo/Dojo 路由工具与动态环境。", "每板块至少 1 题具备真实工具证据"],
        ["P1", "评测环境", "补齐动态靶场、Splunk 与 AgentDojo evaluator。", "人工复核题数量降为 0"],
        ["P2", "成本与稳定性", "记录请求、Token、时延、工具错误与清理率；失败可断点恢复。", "12 题完整且无重复/漏题"],
    ]
    _add_table(document, ["优先级", "工作流", "实施要点", "出口门"], roadmap, [1100, 1600, 4200, 2966], small=True)
    _add_callout(
        document,
        "复测顺序",
        "先跑 1 道静态题验证 final_answer 与证据合同，再跑 1 道动态题验证目标服务和清理，最后顺序执行固定 12 题。",
    )

    _add_heading(document, "附录 A｜实验与评分明细")
    provenance = [
        ["实验 ID", str(data.evaluation.get("experiment_id") or "unavailable")],
        ["报告状态", str(data.evaluation.get("report_status") or "UNKNOWN")],
        ["数据目标", ", ".join(sorted(data.observed_targets)) or "unlabeled"],
        ["模型", f"{data.provider} / {data.model}"],
        ["evaluation.json", str(data.evaluation_path)],
        ["私有答案导出", str(bool(data.evaluation.get("private_answers_exported"))).lower()],
    ]
    _add_table(document, ["字段", "值"], provenance, [2400, 7466], small=True)

    _add_heading(document, "附录 B｜术语")
    glossary = [
        ["确定性评分", "由固定 evaluator 根据结构化结果与私有 oracle 计算；模型不参与评分。"],
        ["目标达成", "题目最终目标被 oracle 验证，而不是 API 调用成功或任务进入终态。"],
        ["MANUAL_REVIEW_REQUIRED", "题目已执行并保留在分母，但当前缺少自动化专用 evaluator。"],
        ["Evidence", "可复核的原始事实、请求响应、路径位置、事件字段或可重放工具产物。"],
        ["结构验证预览", "使用非 PentAGI 数据检查报告生成器，禁止作为 PentAGI 成绩发布。"],
    ]
    _add_table(document, ["术语", "定义"], glossary, [2700, 7166], small=True)

    end = document.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end.add_run("— END OF REPORT —")
    _set_run_font(run, size=8, bold=True, color=MUTED)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Generate a PentAGI deterministic benchmark DOCX report")
    parser.add_argument("--evaluation", type=Path, required=True)
    parser.add_argument("--template-docx", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--comparison-evaluation", type=Path)
    parser.add_argument("--target-name", default="PentAGI")
    parser.add_argument(
        "--preview",
        action="store_true",
        help="Allow non-PentAGI data for structural testing; the report is visibly labeled as a preview",
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    data = load_report_data(args.evaluation, target_name=args.target_name, preview=args.preview)
    comparison = _comparison_summary(args.comparison_evaluation)
    _build_document(
        data,
        template_docx=args.template_docx,
        output_path=args.output.resolve(),
        comparison=comparison,
    )
    print(
        json.dumps(
            {
                "output": str(args.output.resolve()),
                "experiment_id": data.evaluation.get("experiment_id"),
                "target": sorted(data.observed_targets),
                "preview": data.preview,
                "score": data.evaluation.get("equal_weight_category_score"),
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
